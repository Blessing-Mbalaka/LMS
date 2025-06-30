import io
import re
from PyPDF2 import PdfReader
import docx
from docx import Document


def extract_text_from_pdf(f):
    reader = PdfReader(f)
    texts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            texts.append(t)
    return "\n\n".join(texts)

def extract_text_from_docx(f):
    doc = docx.Document(io.BytesIO(f.read()))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

def extract_text(file_obj, content_type):
    """
    f: InMemoryUploadedFile or file-like
    content_type: e.g. 'application/pdf' or docx mime
    """
    file_obj.seek(0)
    if content_type == "application/pdf":
        return extract_text_from_pdf(file_obj)
    elif content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return extract_text_from_docx(file_obj)
    else:
        # fallback: try decoding
        return file_obj.read().decode("utf-8", errors="ignore")


import re

def extract_case_studies(raw_text):
    """
    Scan raw_text for every `Case Study:` block, capturing
    everything until the next question-number marker.
    Returns a dict: { question_number: case_study_text }
    """
    # question numbers look like 1.1 or 1.1.2 etc.
    qn_rx = re.compile(
    r'^\s*(\d+(?:\.\d+)+)\s+(.*?)(?:\s*\((\d+)\s*marks?\))?\s*$', 
    re.IGNORECASE
)


    cs_rx = re.compile(r"Case Study\s*:\s*", re.IGNORECASE)

    lines = raw_text.splitlines()
    cs_map = {}
    current_q = None
    buffer = []

    for line in lines:
        # new question
        m_q = qn_rx.match(line.strip())
        if m_q:
            # if we were buffering a case‐study, attach it to the *previous* qn
            if current_q and buffer:
                cs_map[current_q] = "\n".join(buffer).strip()
            buffer = []
            current_q = m_q.group(1)
            continue

        # case‐study start
        if cs_rx.search(line):
            # begin buffering
            buffer.append(cs_rx.sub("", line).strip())
            continue

        # if buffering, keep adding until next question
        if buffer is not None and buffer != []:
            buffer.append(line.strip())

    # final flush
    if current_q and buffer:
        cs_map[current_q] = "\n".join(buffer).strip()

    return cs_map

#Table extraction regex logic-This will be used to extract tables and questions,
#So what is happening? The code uses a regex delimination approach to identify the 
#XML file components of the word document and use the table fields deliminator to 
# Identify the fields on the table, and populate the data. 

# *In very simple*  terms: The question and its table should get extracted with this table---
#This is going to be an option, called the advanced extraction approach to help 
#structure the data. 

#_____________________________________________________________________________________
#<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<
#<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<
#<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<


#These close imports matter please leave them here...
from docx import Document
from zipfile import BadZipFile


import io
import re
from docx import Document
from zipfile import BadZipFile

# 1. Match Column A with Column B Extractor
def extract_match_column_tables(docx_file):
    try:
        doc = Document(io.BytesIO(docx_file.read()))
    except BadZipFile:
        return {"error": "Uploaded file is not a valid .docx document."}
    except Exception as e:
        return {"error": str(e)}

    match_tables = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)

        if rows and any("Column A" in cell or "Column B" in cell for cell in rows[0]):
            match_tables.append({
                "header": rows[0],
                "rows": rows[1:]
            })

    return match_tables


# 2. Multiple Choice Table Extractor
def extract_multiple_choice_tables(docx_file):
    try:
        doc = Document(io.BytesIO(docx_file.read()))
    except BadZipFile:
        return {"error": "Uploaded file is not a valid .docx document."}
    except Exception as e:
        return {"error": str(e)}

    mcq_tables = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)

        if rows and "Correct Answer" in rows[0]:
            mcq_tables.append({
                "header": rows[0],
                "rows": rows[1:]
            })

    return mcq_tables


# 3. Combined Extractor 
def extract_all_tables(docx_file):
    docx_file.seek(0)
    match_tables = extract_match_column_tables(docx_file)

    docx_file.seek(0)
    mcq_tables = extract_multiple_choice_tables(docx_file)

    docx_file.seek(0)
    generic_tables = extract_generic_tables(docx_file)

    return {
        "match_tables": match_tables,
        "mcq_tables": mcq_tables,
        "generic_tables": generic_tables,
    }


from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from zipfile import BadZipFile
import io
import re

def extract_questions_with_metadata(docx_file):
    docx_file.seek(0)
    try:
        doc = Document(io.BytesIO(docx_file.read()))
    except BadZipFile:
        return {"error": "Uploaded file is not a valid .docx document."}

    # Regex to detect question numbers and marks
    qn_rx = re.compile(
    r'^\s*(\d+(?:\.\d+)+)\s+(.*?)(?:\s*\((\d+)\s*marks?\))?\s*$', 
    re.IGNORECASE
)


    marks_rx = re.compile(r'\((\d+)\s*Marks?\)', re.IGNORECASE)

    results   = {}
    current_q = None
    in_cs     = False   # are we currently inside a case-study?

    for block in doc.element.body:

        # —— Paragraphs ——
        if block.tag.endswith('}p'):
            para = Paragraph(block, doc)
            text = para.text.strip()
            if not text:
                continue

            # 1️⃣ New question header?
            if m := qn_rx.match(text):
                current_q = m.group(1)
                in_cs     = False
                results[current_q] = {
                    "instruction": text,
                    "question_text": "",
                    "case_study": "",
                    "table": [],
                    "marks": marks_rx.search(text).group(1) if marks_rx.search(text) else ""
                }
                continue

            # 2️⃣ Case Study start?
            if current_q and text.lower().startswith("case study"):
                in_cs = True
                # strip off “Case Study:” label
                cs_body = re.sub(r'(?i)^case study\s*:?', '', text).strip()
                results[current_q]["case_study"] += cs_body + "\n"
                continue

            # 3️⃣ Any other paragraph
            if current_q:
                if in_cs:
                    results[current_q]["case_study"] += text + "\n"
                else:
                    results[current_q]["question_text"] += text + "\n"

        # —— Tables ——
        elif block.tag.endswith('}tbl') and current_q:
            tbl = Table(block, doc)
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in tbl.rows
                if any(cell.text.strip() for cell in row.cells)
            ]
            if rows:
                results[current_q]["table"].append(rows)

    # Final trim
    for qdata in results.values():
        qdata["question_text"] = qdata["question_text"].strip()
        qdata["case_study"]    = qdata["case_study"].strip()

    return results





def extract_generic_tables(docx_file):
    try:
        doc = Document(io.BytesIO(docx_file.read()))
    except BadZipFile:
        return {"error": "Uploaded file is not a valid .docx document."}
    except Exception as e:
        return {"error": str(e)}

    generic_tables = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)

        # Skip if it's already handled
        if rows and not ("Column A" in rows[0] or "Correct Answer" in rows[0]):
            generic_tables.append({
                "header": rows[0],
                "rows": rows[1:]
            })

    return generic_tables


#This is the latest view which is meant to extract the full paper as is since we have proved we can extract each component individually.
import io
import re
import json
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from zipfile import BadZipFile
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from google import genai

# Namespaces for inline image extraction
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

@require_POST
@csrf_exempt
def auto_classify_blocks(request):
    """
    Use Gemini API to classify text and figure blocks.
    Returns {"types": [...]} aligned with input.
    """
    try:
        payload = json.loads(request.body)
        blocks  = payload.get('blocks', [])

        system_prompt = """
You are a classification assistant. You receive a JSON list of text/image blocks from an exam paper, in order.
Produce {"types": [...]} matching each block. Valid: question_header, case_study, paragraph, table, instruction, rubric, diagram, figure.

Rules:
0. A hyphen-only block (>=5 '-') is 'instruction' boundary.
1. FIRST block post-boundary starting '1.1' marks questions; pre-block: 'instruction'.
2. Inline image blocks are 'figure'.
3. 'Case Study' start & continuation => 'case_study'.
4. Question headers => 'question_header'.
5. Others post-question split by content; default 'paragraph'.
Output only JSON.
"""
        resp = genai_client.models.generate_content(
            model='gemini-2.0-flash',
            contents=[system_prompt, json.dumps(blocks)]
        )
        data = json.loads((resp.text or '').strip())
        types = data.get('types')
        if not isinstance(types, list) or len(types) != len(blocks):
            raise ValueError
    except Exception:
        types = ['paragraph'] * len(blocks)

    return JsonResponse({'types': types})


import io
import re
import base64
from zipfile import BadZipFile
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

# Namespace constants for inline drawing
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def extract_full_docx_structure(docx_file):
    """
    Extracts a .docx into structured blocks including question headers, paragraphs, 
    tables, images, and case studies. Captures “(10 Marks)” either on the same line
    as the question or as its own paragraph immediately after.
    """
    # Regex definitions
    qn_rx = re.compile(
        r'^\s*(\d+(?:\.\d+)+)\s+(.*?)(?:\(\s*(\d+)\s*marks?\))?\s*$',
        re.IGNORECASE
    )
    marks_only_rx = re.compile(r'^\(?\s*(\d+)\s*marks?\)?$', re.IGNORECASE)
    cs_rx = re.compile(r'(?i)^Case Study\s*:?')

    docx_file.seek(0)
    try:
        doc = Document(io.BytesIO(docx_file.read()))
    except BadZipFile:
        return {'error': 'Invalid .docx'}

    blocks = []
    current_case_study = None

    for el in doc.element.body:
        # -- Paragraph handling --
        if el.tag.endswith('}p'):
            para = Paragraph(el, doc)
            text = para.text.strip()
            if not text:
                continue

            # 1) Standalone marks-only paragraph immediately after a question header
            m_marksonly = marks_only_rx.match(text)
            if m_marksonly and blocks and blocks[-1]['type'] == 'question_header':
                blocks[-1]['marks'] = m_marksonly.group(1)
                continue

            # 2) Inline images
            has_image = False
            for drawing in el.iter(f'{{{W_NS}}}drawing'):
                for blip in drawing.iter(f'{{{A_NS}}}blip'):
                    rid = blip.get(f'{{{R_NS}}}embed')
                    if not rid:
                        continue  # Skip malformed image
                    part = doc.part.related_parts.get(rid)
                    if not part:
                        continue  # Skip broken/missing image
                    b64 = base64.b64encode(part.blob).decode('ascii')
                    uri = f"data:{part.content_type};base64,{b64}"
                    blocks.append({'type': 'figure', 'data_uri': uri})
                    has_image = True
                    current_case_study = None
            if has_image:
                continue  # Don't process this paragraph as text

            # 3) Question header inline
            m_q = qn_rx.match(text)
            if m_q:
                num, rest, mark = m_q.groups()
                blocks.append({
                    'type': 'question_header',
                    'text': f"{num} {rest}".strip(),
                    'marks': mark or ""
                })
                current_case_study = None
                continue

            # 4) Case study start
            if cs_rx.match(text):
                current_case_study = {
                    'type': 'case_study',
                    'text': cs_rx.sub('', text).strip()
                }
                continue

            # 5) Inside case study
            if current_case_study:
                current_case_study['text'] += '\n' + text
                continue

            # 6) Regular paragraph
            blocks.append({'type': 'paragraph', 'text': text})

        # -- Table handling --
        elif el.tag.endswith('}tbl'):
            tbl = Table(el, doc)
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in tbl.rows
                if any(cell.text.strip() for cell in row.cells)
            ]

            # A single-cell table that matches a question header? Treat as header.
            if len(rows) == 1 and len(rows[0]) == 1:
                cell_text = rows[0][0]
                m = qn_rx.match(cell_text)
                if m:
                    num, rest, mark = m.groups()
                    blocks.append({
                        'type': 'question_header',
                        'text': f"{num} {rest}".strip(),
                        'marks': mark or ""
                    })
                    current_case_study = None
                    continue

            # Flush any open case study first
            if current_case_study:
                blocks.append(current_case_study)
                current_case_study = None

            # Otherwise treat as a normal table
            if rows:
                blocks.append({'type': 'table', 'rows': rows})

    # Flush trailing case study
    if current_case_study:
        blocks.append(current_case_study)

    return blocks



import io
import re
import json
import base64
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from zipfile import BadZipFile
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from google import genai

# Namespaces for inline image extraction
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

@require_POST
@csrf_exempt
def auto_classify_blocks(request):
    """
    Calls Gemini to classify each block into one of eight types
    (question_header, case_study, paragraph, table, instruction,
     rubric, diagram, figure) aligned with input.
    """
    try:
        payload = json.loads(request.body)
        blocks  = payload.get('blocks', [])

        # Pre-classify based on simple patterns
        types = []
        for b in blocks:
            text = (b.get('text') or '').strip()
            # Question headers like '1.1', '2.3.4', etc.
            if re.match(r'^\d+(?:\.\d+)+', text):
                types.append('question_header')
                continue
            # Case study markers
            if text.lower().startswith('case study'):
                types.append('case_study')
                continue
            # Inline images
            if b.get('data_uri'):
                types.append('figure')
                continue
            # Tables
            if b.get('rows') is not None:
                types.append('table')
                continue
            # Hyphen separator
            if re.fullmatch(r'-{5,}', text):
                types.append('instruction')
                continue
            # Fallback placeholder, will refine via AI
            types.append(None)

        # Build system prompt for remaining unknowns
        system_prompt = ("""
You are a classification assistant. For any block where I haven't pre-assigned a type,
classify it into one of: paragraph, instruction, rubric, diagram.
Use the context of the content. Do NOT reclassify ones already pre-labeled.
Output JSON only: {"types": [...]} matching each input block.
""")

        # Send blocks and partial-types to Gemini
        resp = genai_client.models.generate_content(
            model='gemini-2.0-flash',
            contents=[system_prompt, json.dumps({'blocks': blocks, 'partial_types': types})]
        )
        data = json.loads((resp.text or '').strip())
        ai_types = data.get('types', [])

        # Merge: use pre-classified if present, else AI result
        final_types = []
        for pre, ai in zip(types, ai_types):
            final_types.append(pre or ai)

    except Exception:
        # On error, default to paragraph
        final_types = ['paragraph'] * len(blocks)

    return JsonResponse({'types': final_types})
